from flask import Flask, render_template, request, jsonify
import os
import re
import docx 
import PyPDF2
from win32com import client as wc
import pythoncom

app = Flask(__name__)

# 获取根目录下的盘符
@app.route('/root-drives', methods=['GET'])
def root_drives():
    drives = []
    for drive in range(ord('A'), ord('Z')+1):
        drive = chr(drive)
        if os.path.exists(drive + ':\\'):
            drives.append(drive)
    return jsonify(drives)

# 获取指定路径下的子文件夹
@app.route('/subdirectories', methods=['GET'])
def subdirectories():
    path = request.args.get('path')
    subdirectories = [subdirectory for subdirectory in os.listdir(path) if os.path.isdir(os.path.join(path, subdirectory))]
    return jsonify(subdirectories)

# 检查文件扩展名是否允许
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ['doc', 'docx', 'pdf']

# 搜索文件
def search_files(path, keyword):
    results = []
    for root, dirs, files in os.walk(path):
        for file in files:
            if allowed_file(file):
                file_path = os.path.join(root, file)
                if search_keyword_in_file(file_path, keyword):
                    result = {
                        'path': file_path,
                        'lines': get_file_content(file_path, keyword)
                    }
                    results.append(result)
    return results

# 检索关键字
def search_keyword_in_file(file_path, keyword):
    pythoncom.CoInitialize()
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if re.search(keyword, paragraph.text, re.IGNORECASE):
                return True
    elif file_path.endswith('.doc'):
        word = wc.Dispatch("Word.Application")
        doc = word.Documents.Open(file_path)
        doc.SaveAs(file_path+'.docx', 12)   
        doc.Close()
        word.Quit()
        doc = docx.Document(file_path+'.docx')
        for paragraph in doc.paragraphs:
            if re.search(keyword, paragraph.text, re.IGNORECASE):
                return True
    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                if re.search(keyword, page.extract_text(), re.IGNORECASE):
                    return True
    return False

# 获取文件内容
def get_file_content(file_path, keyword):
    pythoncom.CoInitialize()
    content = []
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        for i, paragraph in enumerate(doc.paragraphs, start=1):
            if re.search(keyword, paragraph.text, re.IGNORECASE):
                highlighted_text = re.sub(rf'({keyword})', r'<mark>\1</mark>', paragraph.text, flags=re.IGNORECASE)
                content.append({'line': i, 'text': highlighted_text})
    elif file_path.endswith('.doc') :
        # word = wc.Dispatch("Word.Application")
        # doc = word.Documents.Open(file_path)
        # doc.SaveAs(file_path+'.docx', 12)   
        # doc.Close()
        # word.Quit()
        doc = docx.Document(file_path+'.docx')
        for i, paragraph in enumerate(doc.paragraphs, start=1):
            if re.search(keyword, paragraph.text, re.IGNORECASE):
                highlighted_text = re.sub(rf'({keyword})', r'<mark>\1</mark>', paragraph.text, flags=re.IGNORECASE)
                content.append({'line': i, 'text': highlighted_text})
    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                txt = page.extract_text()
                txt=txt.split(" ")
                for i, content_line in enumerate(txt, start=1):
                    if re.search(keyword, content_line, re.IGNORECASE):
                        highlighted_text = re.sub(rf'({keyword})', r'<mark>\1</mark>', content_line, flags=re.IGNORECASE)
                        content.append({'line': i, 'text': highlighted_text})

    return content

# 主页
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        keyword = request.form.get('keyword')
        selected_path = request.form.get('selected_path')
        if selected_path:
            results = search_files(selected_path, keyword)
            return jsonify(results)
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
